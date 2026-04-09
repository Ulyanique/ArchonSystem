from sqlalchemy.ext.asyncio import AsyncSession
from typing import List, Optional
from datetime import datetime
from app.models import Universe, Character, Location, Chapter, Note, TimelineEvent
from app.services import knowledge
from app.utils.author_notes import strip_author_notes
from app.utils.resolve_markers import resolve_markers

class MarkdownExportService:
    """Сервис для экспорта вселенной в Markdown"""
    
    def __init__(self, db: AsyncSession, universe_id: int, master_db: AsyncSession):
        self.db = db
        self.master_db = master_db
        self.universe_id = universe_id
        self.book = None
    
    async def _init_book(self):
        if not self.book:
            self.book = await knowledge.get_universe(self.universe_id, self.master_db)

    async def export_full(self, include_characters: bool = True, include_locations: bool = True,
                    include_chapters: bool = True, include_notes: bool = True,
                    include_timeline: bool = True) -> str:
        """Экспортировать всю вселенную"""
        await self._init_book()
        if not self.book:
            return ""
        
        sections = []
        
        # Заголовок
        sections.append(self._export_title())
        
        # Описание
        if self.book.description:
            sections.append(self._export_description())
        
        # Персонажи
        if include_characters:
            sections.append(await self._export_characters())
        
        # Локации
        if include_locations:
            sections.append(await self._export_locations())
        
        # Главы
        if include_chapters:
            sections.append(await self._export_chapters())
        
        # Заметки
        if include_notes:
            sections.append(await self._export_notes())
        
        # Таймлайн
        if include_timeline:
            sections.append(await self._export_timeline())
        
        return "\n\n".join(sections)
    
    def _export_title(self) -> str:
        """Экспорт заголовка"""
        genre_str = f" ({self.book.genre})" if self.book.genre else ""
        return f"# {self.book.title}{genre_str}\n\n"
    
    def _export_description(self) -> str:
        """Экспорт описания"""
        return f"## Описание\n\n{self.book.description}\n\n"
    
    async def _export_characters(self) -> str:
        """Экспорт персонажей"""
        characters = await knowledge.get_characters(self.db, self.universe_id)
        if not characters:
            return "## Персонажи\n\n_Нет персонажей_\n"
        
        sections = ["## Персонажи\n"]
        for char in characters:
            char_sections = [f"### {char.name}"]
            if char.role: char_sections.append(f"**Роль:** {char.role}")
            if char.traits: char_sections.append(f"**Черты характера:** {char.traits}")
            if char.appearance: char_sections.append(f"**Внешность:** {char.appearance}")
            if char.backstory: char_sections.append(f"**Предыстория:**\n\n{char.backstory}")
            if char.description: char_sections.append(f"**Описание:**\n\n{char.description}")
            sections.append("\n".join(char_sections))
        
        return "\n\n".join(sections)
    
    async def _export_locations(self) -> str:
        """Экспорт локаций"""
        locations = await knowledge.get_locations(self.db, self.universe_id)
        if not locations:
            return "## Локации\n\n_Нет локаций_\n"
        
        sections = ["## Локации\n"]
        for loc in locations:
            loc_sections = [f"### {loc.name}"]
            if loc.location_type: loc_sections.append(f"**Тип:** {loc.location_type}")
            if loc.description: loc_sections.append(f"\n{loc.description}")
            if loc.details: loc_sections.append(f"\n**Детали:**\n\n{loc.details}")
            sections.append("\n".join(loc_sections))
        
        return "\n\n".join(sections)
    
    async def _export_chapters(self) -> str:
        """Экспорт глав"""
        chapters = await knowledge.get_chapters(self.db, self.universe_id)
        if not chapters:
            return "## Главы\n\n_Нет глав_\n"
        
        sections = ["## Главы\n"]
        for chapter in chapters:
            chapter_sections = [f"### Глава {chapter.chapter_number}: {chapter.title}"]
            if chapter.summary: chapter_sections.append(f"\n**Краткое содержание:**\n\n{chapter.summary}\n")
            if chapter.content:
                content = strip_author_notes(chapter.content)
                content = await resolve_markers(self.db, self.universe_id, content)
                chapter_sections.append(f"\n{content}\n")
            if chapter.notes: chapter_sections.append(f"\n---\n*Заметки:* {chapter.notes}\n")
            sections.append("\n".join(chapter_sections))
        
        return "\n\n".join(sections)
    
    async def _export_notes(self) -> str:
        """Экспорт заметок"""
        notes = await knowledge.get_notes(self.db, self.universe_id)
        if not notes:
            return "## Заметки\n\n_Нет заметок_\n"
        
        sections = ["## Заметки\n"]
        type_labels = {'idea': '💡 Идея', 'research': '📚 Исследование', 'draft': '✏️ Черновик', 'other': '📝 Другое'}
        for note in notes:
            type_label = type_labels.get(note.note_type, note.note_type)
            sections.append(f"### {type_label}: {note.title}\n\n{note.content}\n")
        
        return "\n\n".join(sections)
    
    async def _export_timeline(self) -> str:
        """Экспорт таймлайна"""
        from app.services.timeline import get_timeline_events
        events = await get_timeline_events(self.db, self.universe_id)
        if not events:
            return "## Таймлайн\n\n_Нет событий_\n"
        
        sections = ["## Таймлайн событий\n"]
        for event in events:
            event_parts = [f"### {event.title}"]
            if event.date_value: event_parts.append(f"**Время:** {event.date_value}")
            if event.description: event_parts.append(f"\n{event.description}")
            sections.append("\n".join(event_parts))
        
        return "\n\n".join(sections)
    
    async def export_chapter_only(self, chapter_id: int) -> str:
        """Экспортировать только главу"""
        chapter = await knowledge.get_chapter(self.db, chapter_id)
        if not chapter: return ""
        sections = [f"# Глава {chapter.chapter_number}: {chapter.title}\n"]
        if chapter.summary: sections.append(f"## Краткое содержание\n\n{chapter.summary}\n")
        if chapter.content:
            content = strip_author_notes(chapter.content)
            content = await resolve_markers(self.db, self.universe_id, content)
            sections.append(f"\n{content}\n")
        return "\n\n".join(sections)
    
    async def export_characters_only(self) -> str:
        """Экспортировать только персонажей"""
        await self._init_book()
        return await self._export_characters()


class DocxExportService:
    """Сервис для экспорта вселенной в DOCX"""
    
    def __init__(self, db: AsyncSession, universe_id: int, master_db: AsyncSession):
        self.db = db
        self.master_db = master_db
        self.universe_id = universe_id
        self.book = None

    async def _init_book(self):
        if not self.book:
            self.book = await knowledge.get_universe(self.universe_id, self.master_db)
    
    async def export_full(self, include_characters: bool = True, include_locations: bool = True,
                    include_chapters: bool = True, include_notes: bool = True,
                    include_timeline: bool = True):
        """Экспортировать всю вселенную в Document. Возвращает docx.Document."""
        from docx import Document
        await self._init_book()
        doc = Document()
        if not self.book: return doc
        
        title = self.book.title + (f" ({self.book.genre})" if self.book.genre else "")
        doc.add_heading(title, level=0)
        
        if self.book.description:
            doc.add_heading("Описание", level=1)
            doc.add_paragraph(self.book.description)
        
        if include_characters:
            characters = await knowledge.get_characters(self.db, self.universe_id)
            doc.add_heading("Персонажи", level=1)
            if not characters:
                doc.add_paragraph("Нет персонажей")
            else:
                for char in characters:
                    doc.add_heading(char.name, level=2)
                    if char.role: doc.add_paragraph(f"Роль: {char.role}")
                    if char.traits: doc.add_paragraph(f"Черты характера: {char.traits}")
                    if char.appearance: doc.add_paragraph(f"Внешность: {char.appearance}")
                    if char.backstory: doc.add_paragraph(f"Предыстория:\n{char.backstory}")
                    if char.description: doc.add_paragraph(f"Описание:\n{char.description}")
        
        if include_locations:
            locations = await knowledge.get_locations(self.db, self.universe_id)
            doc.add_heading("Локации", level=1)
            if not locations:
                doc.add_paragraph("Нет локаций")
            else:
                for loc in locations:
                    doc.add_heading(loc.name, level=2)
                    if loc.location_type: doc.add_paragraph(f"Тип: {loc.location_type}")
                    if loc.description: doc.add_paragraph(loc.description)
                    if loc.details: doc.add_paragraph(f"Детали:\n{loc.details}")
        
        if include_chapters:
            chapters = await knowledge.get_chapters(self.db, self.universe_id)
            doc.add_heading("Главы", level=1)
            if not chapters:
                doc.add_paragraph("Нет глав")
            else:
                for chapter in chapters:
                    doc.add_heading(f"Глава {chapter.chapter_number}: {chapter.title}", level=2)
                    if chapter.summary: doc.add_paragraph(f"Краткое содержание:\n{chapter.summary}")
                    if chapter.content:
                        content = strip_author_notes(chapter.content)
                        content = await resolve_markers(self.db, self.universe_id, content)
                        doc.add_paragraph(content)
                    if chapter.notes: doc.add_paragraph(f"Заметки: {chapter.notes}")
        
        if include_notes:
            notes = await knowledge.get_notes(self.db, self.universe_id)
            doc.add_heading("Заметки", level=1)
            type_labels = {'idea': 'Идея', 'research': 'Исследование', 'draft': 'Черновик', 'other': 'Другое'}
            if not notes:
                doc.add_paragraph("Нет заметок")
            else:
                for note in notes:
                    type_label = type_labels.get(note.note_type, note.note_type)
                    doc.add_heading(f"{type_label}: {note.title}", level=2)
                    doc.add_paragraph(note.content)
        
        if include_timeline:
            from app.services.timeline import get_timeline_events
            events = await get_timeline_events(self.db, self.universe_id)
            doc.add_heading("Таймлайн событий", level=1)
            if not events:
                doc.add_paragraph("Нет событий")
            else:
                for event in events:
                    doc.add_heading(event.title, level=2)
                    if event.date_value: doc.add_paragraph(f"Время: {event.date_value}")
                    if event.description: doc.add_paragraph(event.description)
        return doc
    
    async def build_bytes_io(self, **kwargs):
        """Собрать DOCX в BytesIO"""
        from io import BytesIO
        doc = await self.export_full(**kwargs)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
